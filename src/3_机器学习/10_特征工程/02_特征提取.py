# -*- encoding: utf-8 -*-
"""
@File    : 02_特征提取.py
@Time    : 2022/4/22 0022 14:31
@Author  : L
@Software: PyCharm
"""

"""
一、字典特征抽取
"""

measurements = [{'city': 'Beijing', 'temperature': 33.}, {'city': 'London', 'temperature': 12.},
                {'city': 'San Fransisco', 'temperature': 18.}]
from sklearn.feature_extraction import DictVectorizer

dv = DictVectorizer()
dv_fit = dv.fit_transform(measurements)
print(dv_fit)

print(dv_fit.toarray())

# 使用get_feature_names()可以查看具体特征名：
print(dv.get_feature_names())
print(dv_fit.toarray())

"""
这本质上使用的是one-hot编码，one hot编码是将类别变量转换为机器学习算法易于利用的一种形式的过程。
简单来说，就是我将所有特征名称列出来，有这一项的标记1，没有就为0。
"""

"""
二、文本特征抽取
"""
from sklearn.feature_extraction.text import CountVectorizer

data = ["i love python,life is so shot", "i dislike python,life is too long"]
cv = CountVectorizer()
data_fit = cv.fit_transform(data)
print(cv.get_feature_names())
print(data_fit.toarray())

"""
三、中文文本特征抽取
"""

from docx import Document

document = Document('guyongzhe.docx')
text = []
for paragraph in document.paragraphs:
    text.append(paragraph.text)

import jieba

data = jieba.cut(str(text))  # 返回的是生成器
data = list(data)  # 转成列表
data = ' '.join(data)  # 用空格分隔

from sklearn.feature_extraction.text import CountVectorizer

cv = CountVectorizer()
cv.fit_transform([data])

from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document
import jieba

document = Document('guyongzhe.docx')
text = []
for paragraph in document.paragraphs:
    text.append(paragraph.text)
data = jieba.cut(str(text))
data = list(data)
data = ' '.join(data)
# cv = CountVectorizer()
# cv.fit_transform([data])
tfidf = TfidfVectorizer()
tf_idf = tfidf.fit_transform([data])
word = tfidf.get_feature_names()
weight = tf_idf.toarray()
for i in range(len(weight)):  # 打印每类文本的tf-idf词语权重，第一个for遍历所有文本，第二个for遍历某一类文本下的词语权重
    for j in range(len(word)):
        print(word[j], weight[i][j])
